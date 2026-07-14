from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# ============================================================
# HALAMAN SAMPUL
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('LAPORAN\nGAME DEVELOPMENT LIFE CYCLE (GDLC)')
run.font.size = Pt(22)
run.font.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Game: SYNERGY (SNERGY)')
run.font.size = Pt(16)
run.font.bold = True
run.font.name = 'Times New Roman'

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('2D Platformer Co-op (Local Multiplayer)')
run.font.size = Pt(13)
run.font.italic = True
run.font.name = 'Times New Roman'

for _ in range(6):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Praktikum Pengembangan Aplikasi Game')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

info2 = doc.add_paragraph()
info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info2.add_run('SYNERGY_DEV')
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

info3 = doc.add_paragraph()
info3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info3.add_run('2026')
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

doc.add_page_break()

# ============================================================
# DAFTAR ISI
# ============================================================
doc.add_heading('DAFTAR ISI', level=1)
toc_items = [
    ('Informasi Umum Proyek', '1'),
    ('Fase 1: Initiation (Inisiasi)', '2'),
    ('   1.1 Latar Belakang', '2'),
    ('   1.2 Konsep Inti Game', '2'),
    ('   1.3 Unique Selling Point (USP)', '2'),
    ('   1.4 Target Audiens', '3'),
    ('Fase 2: Pre-Production (Pra-Produksi)', '3'),
    ('   2.1 Desain Karakter', '3'),
    ('   2.2 Desain Level', '3'),
    ('   2.3 Alur Progresi Level', '4'),
    ('   2.4 Desain Sistem Audio', '4'),
    ('   2.5 Desain Mekanik Gameplay', '5'),
    ('Fase 3: Production (Produksi)', '5'),
    ('   3.1 Arsitektur Skrip', '5'),
    ('   3.2 Implementasi Fitur Utama', '6'),
    ('   3.3 Aset Visual', '9'),
    ('   3.4 Input & Kontrol', '9'),
    ('Fase 4: Testing (Pengujian)', '10'),
    ('   4.1 Pengujian Fungsional', '10'),
    ('   4.2 Pengujian Kompatibilitas', '11'),
    ('   4.3 Bug yang Ditemukan & Diperbaiki', '11'),
    ('Fase 5: Beta', '12'),
    ('Fase 6: Release (Rilis)', '13'),
    ('Lampiran', '14'),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(f'{item}')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(f'\t{page}')
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

doc.add_page_break()

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.line_spacing = 1.5
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

    # Rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

    doc.add_paragraph()
    return table

# ============================================================
# INFORMASI UMUM PROYEK
# ============================================================
add_heading('Informasi Umum Proyek', level=1)

add_table(
    ['Item', 'Detail'],
    [
        ['Nama Game', 'SYNERGY (SNERGY)'],
        ['Genre', '2D Platformer Co-op (Local Multiplayer)'],
        ['Platform Target', 'PC (Windows Standalone)'],
        ['Game Engine', 'Unity 6 (6000.5.3f1)'],
        ['Bahasa Pemrograman', 'C#'],
        ['Resolusi Default', '1920 × 1080'],
        ['Pengembang', 'SYNERGY_DEV'],
        ['Template Proyek', 'Unity 2D Template (com.unity.template.2d@7.0.4)'],
    ]
)

doc.add_page_break()

# ============================================================
# FASE 1: INITIATION
# ============================================================
add_heading('FASE 1: INITIATION (Inisiasi)', level=1)

add_heading('1.1 Latar Belakang', level=2)
add_para('SYNERGY adalah game platformer 2D yang mengusung konsep kerja sama (co-op) lokal antara dua pemain. Nama "SYNERGY" diambil dari filosofi bahwa dua karakter yang saling bekerja sama mampu menyelesaikan tantangan yang tidak bisa ditaklukkan seorang diri. Game ini dikembangkan sebagai bagian dari Praktikum Pengembangan Aplikasi Game.')

add_heading('1.2 Konsep Inti Game', level=2)
add_bullet('Tipe Permainan: Side-scrolling 2D Platformer dengan mekanik kerja sama dua pemain.')
add_bullet('Jumlah Pemain: 2 pemain (lokal — satu keyboard atau dua gamepad).')
add_bullet('Karakter Utama:')
add_bullet('Skeleton — Karakter gesit yang dikendalikan oleh Player 1 (tombol panah / Joystick 1).')
add_bullet('Golem Blue — Karakter kuat dan besar yang dikendalikan oleh Player 2 (WASD / Joystick 2).')
add_bullet('Tujuan Permainan: Kedua pemain harus bekerja sama melewati rintangan, mengalahkan musuh, mengumpulkan koin, dan mencapai portal akhir bersama-sama untuk menyelesaikan setiap level.')

add_heading('1.3 Unique Selling Point (USP)', level=2)
add_bullet('Co-op Wajib: Kedua karakter harus sama-sama memasuki portal akhir agar level dianggap selesai.')
add_bullet('Mekanik Bertumpuk (Character Stacking): Pemain dapat melompat ke atas kepala pemain lainnya dan berjalan bersama-sama sebagai satu kesatuan.')
add_bullet('AI Companion: Jika hanya satu pemain yang bermain, karakter kedua dapat berubah menjadi bot AI pintar yang mengikuti pemain utama secara dinamis (toggle dengan tombol C).')
add_bullet('Dukungan Dual Gamepad: Dua joystick independen untuk pengalaman co-op yang lebih immersive.')

add_heading('1.4 Target Audiens', level=2)
add_bullet('Mahasiswa dan pelajar penggemar game platformer retro.')
add_bullet('Pemain kasual yang menikmati permainan bersama teman secara lokal.')
add_bullet('Usia 10 tahun ke atas.')

doc.add_page_break()

# ============================================================
# FASE 2: PRE-PRODUCTION
# ============================================================
add_heading('FASE 2: PRE-PRODUCTION (Pra-Produksi)', level=1)

add_heading('2.1 Desain Karakter', level=2)
add_table(
    ['Karakter', 'Peran', 'Kontrol Keyboard', 'Kontrol Gamepad', 'Kecepatan', 'Lompat'],
    [
        ['Skeleton', 'Player 1', 'Panah + Space', 'Joystick 1 + Tombol A', '3', '7'],
        ['Golem Blue', 'Player 2', 'WASD + W (jump)', 'Joystick 2 + Tombol A', '4', '7'],
    ]
)

add_heading('2.2 Desain Level', level=2)
add_para('Game SNERGY memiliki 5 scene utama yang terdaftar dalam Build Settings:')
add_table(
    ['Index', 'Nama Scene', 'Deskripsi', 'Ukuran File'],
    [
        ['0', 'MainMenu', 'Halaman utama permainan dengan tombol Play Game dan Exit', '6 KB'],
        ['1', 'SampleScene', 'Level percobaan/awal (level pertama yang dimainkan)', '323 KB'],
        ['2', 'World_1_Tutorial', 'Level tutorial dengan rintangan dasar', '569 KB'],
        ['3', 'World_2_Forest', 'Level hutan dengan tingkat kesulitan menengah', '210 KB'],
        ['4', 'World_3_Summit', 'Level puncak sebagai tantangan akhir', '210 KB'],
    ]
)

add_heading('2.3 Alur Progresi Level (Game Flow)', level=2)
add_para('Alur progresi level dalam game SNERGY adalah sebagai berikut:')
add_para('MainMenu → SampleScene → World_1_Tutorial → World_2_Forest → World_3_Summit → Kembali ke MainMenu', bold=True)
add_para('Jika pemain mengalami Game Over di level manapun, game akan otomatis kembali ke MainMenu.')

add_heading('2.4 Desain Sistem Audio', level=2)
add_table(
    ['Nama File Audio', 'Format', 'Ukuran', 'Fungsi'],
    [
        ['musik_main_menu.wav', 'WAV', '22.6 MB', 'BGM Main Menu & Gameplay (looping)'],
        ['jump.mp3', 'MP3', '9 KB', 'SFX saat karakter melompat'],
        ['hit.mp3', 'MP3', '68 KB', 'SFX saat karakter terkena serangan musuh'],
        ['coin.mp3', 'MP3', '36 KB', 'SFX saat karakter mengambil koin'],
        ['win.mp3', 'MP3', '89 KB', 'Musik kemenangan saat menyelesaikan level'],
        ['game_over.mp3', 'MP3', '124 KB', 'SFX game over (diputar sekali, tidak looping)'],
        ['dead.mp3', 'MP3', '98 KB', 'SFX kematian karakter'],
    ]
)

add_heading('2.5 Desain Mekanik Gameplay', level=2)
add_para('Mekanik gameplay yang dirancang dalam game SNERGY meliputi:', bold=True)

add_para('A. Pergerakan:', bold=True)
add_bullet('Berjalan ke kiri dan kanan.')
add_bullet('Melompat (dengan deteksi ground).')
add_bullet('Naik tangga (mekanik panjat).')

add_para('B. Pertarungan:', bold=True)
add_bullet('Knockback saat terkena serangan musuh.')
add_bullet('Sistem darah (HP = 100) yang berkurang saat terkena serangan.')
add_bullet('Animasi Hit dan Dead pada karakter.')

add_para('C. Kerja Sama:', bold=True)
add_bullet('Kedua pemain harus masuk portal akhir bersamaan.')
add_bullet('Character Stacking (bertumpuk di atas karakter lain).')
add_bullet('AI Companion yang dapat di-toggle dengan tombol C.')
add_bullet('Batas jarak antar pemain agar tidak terlalu berjauhan.')

add_para('D. Collectible:', bold=True)
add_bullet('Koin dengan skor terpisah per karakter.')
add_bullet('Efek partikel emas saat mengambil koin.')
add_bullet('Skor ditampilkan di layar akhir (victory screen).')

doc.add_page_break()

# ============================================================
# FASE 3: PRODUCTION
# ============================================================
add_heading('FASE 3: PRODUCTION (Produksi)', level=1)

add_heading('3.1 Arsitektur Skrip (Script Architecture)', level=2)
add_para('Proyek ini menggunakan arsitektur berbasis komponen (Component-Based Architecture) khas Unity, dengan pemisahan tanggung jawab ke dalam folder-folder berikut:')

add_table(
    ['No', 'Nama File', 'Lokasi', 'Fungsi Utama'],
    [
        ['1', 'PlayerMovement.cs', 'Scripts/Character/', 'Kontrol pergerakan, lompat, HP, knockback, HUD Skeleton'],
        ['2', 'GolemBlue.cs', 'Scripts/Character/', 'Kontrol pergerakan, lompat, HP, knockback, HUD Golem'],
        ['3', 'CompanionAI.cs', 'Scripts/Character/', 'Sistem AI bot pendamping (toggle C)'],
        ['4', 'CharacterBase.cs', 'Scripts/Character/', 'Kelas dasar karakter (abstraksi)'],
        ['5', 'Aqua.cs', 'Scripts/Character/', 'Karakter elemen air'],
        ['6', 'Pyro.cs', 'Scripts/Character/', 'Karakter elemen api'],
        ['7', 'Terra.cs', 'Scripts/Character/', 'Karakter elemen tanah'],
        ['8', 'Windy.cs', 'Scripts/Character/', 'Karakter elemen angin'],
        ['9', 'EnemyFollow.cs', 'Scripts/Enemy/', 'AI musuh: deteksi, kejar, serang'],
        ['10', 'GameManager.cs', 'Scripts/Managers/', 'Singleton: skor koin, restart, game over'],
        ['11', 'AudioManager.cs', 'Scripts/Managers/', 'Singleton: pemuatan & pemutaran audio'],
        ['12', 'Coin.cs', 'Scripts/Objects/', 'Koin collectible + partikel + SFX'],
        ['13', 'FinishPortal.cs', 'Scripts/Objects/', 'Portal akhir + layar kemenangan'],
        ['14', 'Killzone.cs', 'Scripts/Objects/', 'Area kematian (jurang)'],
        ['15', 'Ladder.cs', 'Scripts/Objects/', 'Mekanik panjat tangga'],
        ['16', 'LightTorch.cs', 'Scripts/Objects/', 'Efek cahaya obor'],
        ['17', 'StartPoint.cs', 'Scripts/Objects/', 'Titik spawn awal pemain'],
        ['18', 'MainMenu.cs', 'Scripts/', 'Logika halaman utama'],
        ['19', 'MenuGameManager.cs', 'Scripts/', 'Overlay menu dalam game'],
        ['20', 'SwapSystem.cs', 'Scripts/Systems/', 'Sistem pergantian karakter'],
        ['21', 'InputManager.cs', 'Scripts/Systems/', 'Pengelolaan input'],
        ['22', 'TileSystem.cs', 'Scripts/Systems/', 'Sistem tile/peta'],
    ]
)

add_heading('3.2 Implementasi Fitur Utama', level=2)

add_heading('3.2.1 Sistem Pergerakan Karakter', level=3)
add_para('Kedua karakter (PlayerMovement.cs dan GolemBlue.cs) mengimplementasikan fitur:')
add_bullet('Pergerakan horizontal (kiri/kanan) menggunakan Rigidbody2D.linearVelocity.')
add_bullet('Lompat dengan deteksi ground (isGrounded) melalui Raycast atau Collision.')
add_bullet('Sistem knockback saat terkena serangan musuh (knockbackForceX, knockbackForceY, knockbackDuration).')
add_bullet('Sistem darah (health = 100) yang berkurang saat terkena serangan.')
add_bullet('Animasi: IsRunning, IsJumping, Hit, Dead melalui Animator Controller.')

add_heading('3.2.2 Sistem AI Musuh (EnemyFollow.cs)', level=3)
add_bullet('Deteksi Pemain Terdekat: Menggunakan FindGameObjectsWithTag("Player") dengan interval optimasi (targetSearchInterval = 0.2s).')
add_bullet('Mengejar & Menyerang: Musuh bergerak ke arah pemain terdekat jika dalam detectionRange, dan menyerang jika dalam attackRange.')
add_bullet('Dukungan Multi-Target: Musuh dapat menyerang baik Skeleton (PlayerMovement) maupun Golem (GolemBlue).')

add_heading('3.2.3 Sistem Audio Terpusat (AudioManager.cs)', level=3)
add_bullet('Pola Singleton Dinamis: Tercipta otomatis saat runtime tanpa perlu dipasang manual di scene.')
add_bullet('Pemuatan Dinamis: Semua klip audio dimuat dari folder Resources/Music menggunakan Resources.Load<AudioClip>().')
add_bullet('Dua AudioSource: Satu untuk musik latar (musicSource, looping) dan satu untuk efek suara (sfxSource, one-shot).')
add_bullet('Metode Publik: PlayMainMenuMusic(), PlayWinMusic(), PlayJumpSFX(), PlayHitSFX(), PlayCoinSFX(), PlayGameOverSFX().')

add_heading('3.2.4 Sistem Koin & Skor (Coin.cs + GameManager.cs)', level=3)
add_bullet('Deteksi Pengambil Koin: Skrip Coin.cs memeriksa apakah yang menyentuh koin memiliki komponen PlayerMovement (Skeleton) atau GolemBlue (Golem), lalu menambahkan skor ke karakter yang sesuai.')
add_bullet('Efek Visual Partikel: Saat koin diambil, sistem membuat ParticleSystem secara dinamis via kode — burst 15 partikel emas yang memudar (fade out) dari posisi koin.')
add_bullet('HUD Skor: Ditampilkan di bawah HP masing-masing karakter (pojok kiri atas untuk Skeleton, pojok kanan atas untuk Golem).')
add_bullet('Rekap Akhir Level: Skor koin masing-masing karakter ditampilkan di panel layar kemenangan (FinishPortal.cs).')

add_heading('3.2.5 Mekanik Bertumpuk (Character Stacking)', level=3)
add_bullet('Pemain dapat melompat ke atas kepala pemain lainnya.')
add_bullet('Karakter di bawah (carrier) mendapatkan pengurangan kecepatan gerak sebesar 40% (carryingSpeedMultiplier = 0.6).')
add_bullet('Karakter di atas (passenger) ikut terbawa bergerak bersama tanpa tergelincir.')

add_heading('3.2.6 AI Companion (CompanionAI.cs)', level=3)
add_bullet('Karakter kedua dapat diubah menjadi bot AI pintar dengan menekan tombol C.')
add_bullet('Fitur AI: mengejar pemain utama, auto-jump rintangan, auto-jump jurang, lompat vertikal.')
add_bullet('Indikator visual teks "AI COMPANION" berwarna cyan di atas kepala bot.')

add_heading('3.2.7 Portal Akhir & Layar Kemenangan (FinishPortal.cs)', level=3)
add_bullet('Kedua pemain harus masuk portal untuk menyelesaikan level.')
add_bullet('Musik kemenangan (win.mp3) dimainkan.')
add_bullet('Layar hasil menampilkan: Judul "LEVEL CLEARED!", Skor koin Skeleton dan Golem, Tombol Next Level, Restart Level, dan Main Menu.')

add_heading('3.2.8 Mekanisme Game Over', level=3)
add_bullet('Ketika health <= 0, karakter memutar animasi mati dan SFX game over sekali saja (tanpa loop).')
add_bullet('Setelah animasi selesai, Time.timeScale dipulihkan ke 1f dan game otomatis memuat scene MainMenu.')

add_heading('3.3 Aset Visual yang Digunakan', level=2)
add_table(
    ['Kategori', 'Sumber Aset'],
    [
        ['Karakter Pemain', 'Sprites/Characters/Player/'],
        ['Musuh', 'Sprites/Characters/Enemy/, Enemy_Galore_I/, DarkFantasyEnemies_FREE/'],
        ['Tileset Lingkungan', 'Sprites/Environment/Nature Platformer Tileset/, Mossy Assets/'],
        ['Objek Desa (Portal, Peti)', 'Cainos/Pixel Art Platformer - Village Props/'],
        ['Item Collectible', 'Pixel Adventure 1/Assets/Items/Fruits/'],
        ['Animasi Tumbuhan', 'Sprites/Environment/Plant Animations 1/'],
    ]
)

add_heading('3.4 Input & Kontrol', level=2)
add_table(
    ['Aksi', 'Player 1 (Skeleton)', 'Player 2 (Golem)'],
    [
        ['Gerak Kiri/Kanan', 'Panah Kiri/Kanan', 'A / D'],
        ['Lompat', 'Space', 'W'],
        ['Toggle AI Companion', 'C', 'C'],
        ['Gamepad Gerak', 'Joystick 1 Analog Kiri', 'Joystick 2 Analog Kiri'],
        ['Gamepad Lompat', 'Joystick 1 Tombol A/B', 'Joystick 2 Tombol A/B'],
    ]
)

doc.add_page_break()

# ============================================================
# FASE 4: TESTING
# ============================================================
add_heading('FASE 4: TESTING (Pengujian)', level=1)

add_heading('4.1 Pengujian Fungsional (Functional Testing)', level=2)
add_table(
    ['No', 'Fitur yang Diuji', 'Skenario Pengujian', 'Hasil yang Diharapkan', 'Status'],
    [
        ['1', 'Pergerakan Skeleton', 'Tekan tombol panah kiri/kanan', 'Skeleton bergerak ke arah yang sesuai', 'Berhasil'],
        ['2', 'Pergerakan Golem', 'Tekan tombol A/D', 'Golem bergerak ke arah yang sesuai', 'Berhasil'],
        ['3', 'Lompat Skeleton', 'Tekan Space saat di tanah', 'Skeleton melompat ke atas', 'Berhasil'],
        ['4', 'Lompat Golem', 'Tekan W saat di tanah', 'Golem melompat ke atas', 'Berhasil'],
        ['5', 'Knockback dari Musuh', 'Karakter mengenai musuh', 'Terpental, HP berkurang, SFX hit berbunyi', 'Berhasil'],
        ['6', 'Pengambilan Koin', 'Karakter menyentuh koin', 'Skor +1, partikel muncul, SFX coin, koin hilang', 'Berhasil'],
        ['7', 'Skor Per Karakter', 'Skeleton ambil koin, Golem ambil koin lain', 'Skor masing-masing bertambah sesuai pengambil', 'Berhasil'],
        ['8', 'Portal Akhir (1 pemain)', 'Hanya 1 pemain masuk portal', 'Level belum selesai, menunggu pemain kedua', 'Berhasil'],
        ['9', 'Portal Akhir (2 pemain)', 'Kedua pemain masuk portal', 'Layar kemenangan muncul dengan skor & tombol', 'Berhasil'],
        ['10', 'Game Over', 'HP karakter habis', 'SFX game over 1x, animasi mati, kembali ke MainMenu', 'Berhasil'],
        ['11', 'Killzone', 'Karakter jatuh ke bawah level', 'Scene aktif dimuat ulang', 'Berhasil'],
        ['12', 'Main Menu → Play', 'Klik tombol Play Game', 'Memuat scene SampleScene', 'Berhasil'],
        ['13', 'Main Menu → Exit', 'Klik tombol Exit Game', 'Aplikasi keluar', 'Berhasil'],
        ['14', 'BGM Main Menu', 'Masuk ke scene MainMenu', 'Musik main menu berputar looping', 'Berhasil'],
        ['15', 'Progresi Level', 'Selesaikan SampleScene, klik Next Level', 'Memuat World_1_Tutorial', 'Berhasil'],
        ['16', 'AI Companion', 'Tekan C saat gameplay', 'Karakter kedua menjadi bot AI', 'Berhasil'],
        ['17', 'Character Stacking', 'Lompat ke atas kepala pemain lain', 'Karakter atas ikut terbawa, carrier melambat', 'Berhasil'],
        ['18', 'Dual Gamepad', 'Sambungkan 2 gamepad, gerakkan', 'Masing-masing joystick mengendalikan karakter berbeda', 'Berhasil'],
    ]
)

add_heading('4.2 Pengujian Kompatibilitas', level=2)
add_table(
    ['No', 'Aspek', 'Detail', 'Status'],
    [
        ['1', 'Unity 6 Compatibility', 'API diperbarui ke FindAnyObjectByType (bukan FindFirstObjectByType yang obsolete)', 'Bersih'],
        ['2', 'Compiler Warnings', 'Peringatan CS0618 telah diatasi di AudioManager.cs dan FinishPortal.cs', 'Bersih'],
        ['3', 'Editor vs Build', 'Scene MainMenu (index 0) berjalan benar di Editor maupun build .exe', 'Berhasil'],
    ]
)

add_heading('4.3 Bug yang Ditemukan & Diperbaiki', level=2)
add_table(
    ['No', 'Bug', 'Penyebab', 'Solusi'],
    [
        ['1', 'Game freeze saat game over, suara loop tanpa henti', 'Time.timeScale = 0 menghentikan coroutine, audio looping', 'Memutar SFX game over 1x (non-loop), memulihkan timeScale, lalu memuat MainMenu'],
        ['2', 'Musik di Main Menu adalah suara game over', 'File musik_main_menu.mp3 adalah duplikat rusak', 'Menggunakan file asli musik_main_menu.wav (22.6 MB)'],
        ['3', 'Error kompilasi TreeViewState is obsolete', 'Skrip editor pihak ketiga menggunakan API lama', 'Membungkus isi file dengan #if false ... #endif'],
        ['4', 'Efek suara hit tidak berbunyi', 'File hit.mp3 belum tersalin ke Resources/Music', 'Menyalin hit.mp3 ke Assets/Resources/Music/'],
        ['5', 'Error partikel koin', 'ParticleSystem otomatis bermain saat AddComponent', 'Menambahkan ps.Stop() sebelum konfigurasi parameter'],
    ]
)

doc.add_page_break()

# ============================================================
# FASE 5: BETA
# ============================================================
add_heading('FASE 5: BETA', level=1)

add_heading('5.1 Status Beta', level=2)
add_para('Game saat ini berada dalam tahap Beta Internal. Semua fitur inti telah diimplementasikan dan diuji secara fungsional di Unity Editor.')

add_heading('5.2 Fitur yang Sudah Stabil', level=2)
fitur_stabil = [
    'Pergerakan dua karakter (Keyboard & Gamepad)',
    'Sistem HP dan Knockback',
    'AI Musuh (mengejar dan menyerang)',
    'AI Companion (bot pendamping)',
    'Character Stacking (mekanik bertumpuk)',
    'Sistem koin dengan skor per karakter',
    'Efek partikel pengambilan koin',
    'Sistem audio terintegrasi (BGM + 7 SFX)',
    'Main Menu dengan Play/Exit',
    'Layar kemenangan dengan rekap skor',
    'Progresi 4 level + Main Menu',
    'Mekanisme game over otomatis kembali ke Main Menu',
    'Killzone di semua level',
]
for fitur in fitur_stabil:
    add_bullet(f'[✓] {fitur}')

add_heading('5.3 Area yang Masih Perlu Pengembangan', level=2)
area_dev = [
    'UI/HUD modern (Canvas-based Health Bar & Coin Icon)',
    'Rintangan co-op sejati (Pressure Plate, mekanik melempar)',
    'Patroli AI musuh saat idle',
    'Sistem penyimpanan progres (Level Select + PlayerPrefs)',
    'Menu pengaturan volume suara (Settings)',
    'Optimasi performa untuk build produksi',
]
for area in area_dev:
    add_bullet(f'[ ] {area}')

doc.add_page_break()

# ============================================================
# FASE 6: RELEASE
# ============================================================
add_heading('FASE 6: RELEASE (Rilis)', level=1)

add_heading('6.1 Build Settings', level=2)
add_table(
    ['Index', 'Scene', 'Keterangan'],
    [
        ['0', 'MainMenu', 'Halaman pertama saat game dimulai'],
        ['1', 'SampleScene', 'Level awal'],
        ['2', 'World_1_Tutorial', 'Level 1 - Tutorial'],
        ['3', 'World_2_Forest', 'Level 2 - Hutan'],
        ['4', 'World_3_Summit', 'Level 3 - Puncak'],
    ]
)
add_para('Target Platform: Windows Standalone')
add_para('Company Name: SYNERGY_DEV')
add_para('Product Name: SYNERGY')

add_heading('6.2 Cara Melakukan Build', level=2)
add_bullet('Buka Unity Editor.')
add_bullet('Pilih menu File > Build Settings...')
add_bullet('Pastikan semua 5 scene terdaftar dengan urutan yang benar.')
add_bullet('Pilih platform PC, Mac & Linux Standalone.')
add_bullet('Klik Build dan tentukan folder output.')
add_bullet('Jalankan file .exe yang dihasilkan.')

add_heading('6.3 Spesifikasi Minimum yang Direkomendasikan', level=2)
add_table(
    ['Komponen', 'Spesifikasi'],
    [
        ['OS', 'Windows 10 / 11'],
        ['Processor', 'Intel Core i3 atau setara'],
        ['RAM', '4 GB'],
        ['GPU', 'Mendukung DirectX 11'],
        ['Storage', '100 MB ruang kosong'],
        ['Input', 'Keyboard + Mouse, atau 2 Gamepad USB'],
    ]
)

doc.add_page_break()

# ============================================================
# LAMPIRAN
# ============================================================
add_heading('LAMPIRAN', level=1)

add_heading('Daftar Seluruh Aset Audio', level=2)
add_table(
    ['No', 'Nama File', 'Format', 'Ukuran', 'Fungsi'],
    [
        ['1', 'musik_main_menu.wav', 'WAV', '22.6 MB', 'BGM Main Menu & Gameplay'],
        ['2', 'jump.mp3', 'MP3', '9 KB', 'SFX lompat'],
        ['3', 'hit.mp3', 'MP3', '68 KB', 'SFX terkena serangan'],
        ['4', 'coin.mp3', 'MP3', '36 KB', 'SFX mengambil koin'],
        ['5', 'win.mp3', 'MP3', '89 KB', 'Musik kemenangan'],
        ['6', 'game_over.mp3', 'MP3', '124 KB', 'SFX game over'],
        ['7', 'dead.mp3', 'MP3', '98 KB', 'SFX kematian'],
    ]
)

add_heading('Struktur Folder Proyek', level=2)
add_para('Assets/', bold=True)
add_bullet('Audio/Music/ — Berkas audio asli (sumber)')
add_bullet('Resources/Music/ — Berkas audio untuk pemuatan dinamis (Resources.Load)')
add_bullet('Scenes/ — File scene Unity (MainMenu, SampleScene, World_1-3)')
add_bullet('Scripts/ — Seluruh skrip C# (Character, Enemy, Managers, Objects, Systems)')
add_bullet('Sprites/ — Sprite karakter, musuh, lingkungan')
add_bullet('Cainos/ — Asset pack Village Props (portal, peti emas)')
add_bullet('Pixel Adventure 1/ — Asset pack Items (buah/koin)')

# Save
output_path = r'd:\PENGEMBANGAN APLIKASI GAME\Praktikum\SNERGY\Laporan_GDLC_SNERGY.docx'
doc.save(output_path)
print(f'Laporan berhasil disimpan di: {output_path}')
